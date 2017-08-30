# -*- coding: utf-8 -*-
"""
Created on Sat Apr  1 20:31:57 2017
"""

#import pandas as pd
import pickle
import re
import json
import sys,getopt
#from docx import Document
from collections import OrderedDict

import docx
from docx import Document
from doc_XML import docx_to_text
import pdfMiner




class document(object):
    """Base class for all types of documents"""

    def __init__(self, text = '', path = '', option = 'raw_text', objects = None):
        """Create text based on document extension, classify text """
        self.objects = objects
        self.section_positions = None
        if option == 'raw_text':
            self.text = str(text).lower()
            self.paragraphs = self.text.split('\n')
        elif option == 'docx':
            self.document=Document(path)
            try:
                self.text = docx_to_text(self.document)
            except:
                try:
                    self.paragraphs= list(map(lambda x: x.text,self.document.paragraphs))
                    self.text='\n'.join(self.paragraphs)
                except:
                    self.text = ' '
            self.objects = self.document
        elif option == 'pdf':
            try:
                text, objects = pdfMiner.to_text(path)
                self.text = str(text).lower()
                self.paragraphs = self.text.split('\n')
                self.objects = objects
                sorted_par = pdfMiner.sort_layout(self.objects)
                out = []
                for item in sorted_par.values():
                    tup = []
                    for more in item:
                        tup.append((more[0].split('\n')[0],more[1]))
                    tup = sorted(tup,key = lambda x: x[1])
                    out.append(list(list(zip(*tup))[0]))
                    self.paragraphs = list(map(lambda item: ' '.join(item),out))
                    self.text = '\n'.join(self.paragraphs)
            except:
                self.text = ' '
        self.classification = self.classify_text()

    def print_info(self):
        self.get_data()
        self.get_pos()
        print(self.section_positions)

    def get_pos(self):
        """Relative appearance of sections(e.g. education, personal info,
        skills)"""
        if len(self.section_positions):
            for i,section in enumerate(self.section_positions.items()):
                index=section[1]['index']
                if index == -1:
                    pass
                else:
                    try:
                        section[1]['text']= self.paragraphs[index : list(self.section_positions.items())[i+1][1]['index']]
                    except:
                        section[1]['text'] = self.paragraphs[index : len(self.paragraphs)]

    def classify_text(self):
        """Opens previously trained tfidf vectorizer and MNB classifier"""
        with open('my_dumped_classifier.pkl', 'rb') as fid:
            classifier_loaded = pickle.load(fid)
        with open('tfidf.pkl', 'rb') as fid:
            vectorizer = pickle.load(fid)

        tested_case = vectorizer.transform([self.text])
        pred = classifier_loaded.predict(tested_case)
        return pred, self.text, self.objects

    def get_data(self):
        """General regexes (e.g. address pattern matching, telephone, mail)"""
        self.get_pos()
        self.mail = re.findall('[\w\.-]+@[\w\.-]+', self.text)
        self.phone = re.findall('(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})'
                                  ,self.text)
        self.date = []
        self.date.append( re.findall(r'(?:(?<=date).*)',self.text))
        self.date.append( re.findall(r'(\d{2})[/.-](\d{2})[/.-](\d{4})$',self.text))
        self.date.append(re.findall( r'''(january.*|february.*|march.*|april.*
                                          |may.*|june.*|july.*|
        august.*|september.*|october.*|november.*|december.*|jan.*|feb.*
        |mar .*|apr.*|aug .*|sept.*|oct .*|nov .*|dec .*).*[0-9]*.*''',self.text))
        self.address = re.findall(r'(?:(?<=address).*[0-9].*)',self.text)
        self.name = re.findall(r'(?:(?<=name).*[\n])','\n'.join(self.section_positions['Intro Info']['text']))
        if len(self.address) == 0:
           self.address = re.findall(r"([0-9][0-9]+[^0-9]*[a-z]*[\r])",self.text)
        if len(self.name) == 0:
            try:
                firstLine=self.text.split('\n')[0]
                secondLine = self.text.split('\n')[1]
            except:
                firstLine = ''
                secondLine = ''
            if (not 'cv' in firstLine and not 'curriculum' in firstLine) \
            and not any(char.isdigit() for char in firstLine):
                self.name = firstLine
            else:
                self.name = secondLine
        self.positions = dict(self.section_positions)
        self.positions['Mail'] = self.mail
        self.positions['Phone'] = self.phone
        self.positions['Date']  = self.date
        self.positions['Address'] = self.address
        self.positions['Name'] = self.name
        return self.positions,self.text

    def to_json(self,file_name):
        with open(file_name + ".json", "w") as file:
            json.dump(self.positions,file, indent=4)

#    @property
#    def positions(self):
#        return self.positions


class invoice(document):

    def __init__(self, text = '', path = '', option = 'raw_text',objects=None):
        super().__init__(text=text,path=path,option=option,objects=objects)
        self.section_positions= {'Intro Info' : {'index' :0, 'text' : ''},
                                 'Money' : {'index' :-1, 'text' : ''} ,
                                 'Other' : {'index' :-1, 'text' : ''}}
        for i,paragraph in enumerate(self.paragraphs):
            paragraph = str(paragraph).lower()
            if 'description' in paragraph or 'quantity' in  paragraph :
                self.section_positions['Money'] = {'index' : i, 'text' : ''}
            if 'total' in paragraph :
                    try:
                        self.section_positions['Other'] = {'index' : i+1, 'text' : ''}
                    except:
                        self.section_positions['Other'] = {'index' : i, 'text' : ''}
        self.section_positions = OrderedDict(sorted(self.section_positions.items(),
                                                    key=lambda t: t[1]['index']))

    def print_info(self):
        self.get_pos()
        print(self.section_positions)

    def get_data(self):
        self.get_pos()
        self.positions, self.text = super().get_data()
        self.positions['Invoice number'] = re.findall(r"(?:(?<=number)|(?<=#)|(?<=no)).[0-9]*",self.text )
        self.positions['Total'] = re.findall(r"(?:(?<!subtotal)(?<!sub total)(?<=total).*[0-9]*)", self.text)
        self.positions['Subtotal'] = re.findall(r"(?:(?<=subtotal)|(?<=sub total).*[0-9]*)", self.text)
        return self.positions,self.text



class letter(document):

    def __init__(self, text = '', path = '', option = 'raw_text', objects = None):

        super().__init__(text=text,path=path,option=option,objects=objects)
        self.section_positions = {'Intro Info' : {'index' :0, 'text' : ''},
                                 'Body' : {'index' :-1, 'text' : ''} ,
                                 'Author' : {'index' :-1, 'text' : ''} }
        for i,paragraph in enumerate(self.paragraphs):
            paragraph=str(paragraph).lower()

            if 'dear' in  paragraph  or 'to whom' in paragraph:
                self.section_positions['Body'] = {'index' : i, 'text' : ''}
            if ('sincerely' in paragraph or 'yours' in paragraph
                or 'faithfully' in paragraph or 'regards' in paragraph
                or 'salutations' in paragraph or 'looking forwards' in paragraph):
                self.section_positions['Author'] = {'index' : i, 'text' : ''}
        self.section_positions = OrderedDict(sorted(self.section_positions.items(), key=lambda t: t[1]['index']))

    def print_info(self):
        self.get_pos()
        print(self.section_positions)

    def get_data(self):
        self.get_pos()
        self.positions, self.text = super().get_data()
        return self.positions, self.text


class cv(document):

    def __init__(self, text = '', path = '', option = 'raw_text', objects = None):

        super().__init__(text=text,path=path,option=option,objects = objects)

        self.section_positions= {'Intro Info' : {'index' :0, 'text' : ''},
                                 'Education' : {'index' :-1, 'text' : ''} ,
                                 'Experience' : {'index' :-1, 'text' : ''} ,
                                 'Language' : {'index' :-1, 'text' : ''},
                                 'Skills' : {'index' :-1, 'text' : ''},
                                 'Closing': {'index' :-1, 'text' : ''}}
        language_flag = False
        for i,paragraph in enumerate(self.paragraphs):
            paragraph=str(paragraph).lower()
            if 'education' in  paragraph :
                self.section_positions['Education'] = {'index' : i, 'text' : ''}
            if 'experience' in paragraph :
                self.section_positions['Experience']= {'index' : i, 'text' : ''}
            if 'languages' in paragraph and not language_flag :
                self.section_positions['Language'] = {'index' : i, 'text' : ''}
                language_flag=True
            if 'skill' in paragraph :
                self.section_positions['Skills'] = {'index' : i, 'text' : ''}
            if 'hereby' in paragraph :
                self.section_positions['Closing'] = {'index' : i, 'text' : ''}
        self.section_positions = OrderedDict(sorted(self.section_positions.items(), key=lambda t: t[1]['index']))

    def print_info(self):
        self.get_pos()
        print(self.section_positions)

    def get_data(self):
        self.get_pos()
        self.positions, self.text = super().get_data()
        return self.positions, self.text

def main(argv):
""" Run classifier from command line with different options, to be invoked by
PHP server side app"""
    #Raw text simulation
#    df = pd.read_csv(r'path')
#    df = df.drop('Unnamed: 0',axis=1)
#    df = df.dropna()
#    in_text = df.ix[12,'text']
    inputfile = ''
    outputfile = ''
    option = ''
    try:
      opts, args = getopt.getopt(argv,"hi:o:t:",["ifile=","ofile=","type="])
    except getopt.GetoptError:
      print ('USAGE : docx_test2.py -i <inputfile> -o <outputfile> -t <document type>')
      sys.exit(2)
    for opt, arg in opts:
      if opt == '-h':
         print ('USAGE : docx_test2.py -i <inputfile> -o <outputfile> -t <document type>')
         sys.exit()
      elif opt in ("-i", "--ifile"):
         inputfile = arg
         print (inputfile)
      elif opt in ("-o", "--ofile"):
         outputfile = arg
         print (inputfile)
      elif opt in ("-t", "--type"):
         option = arg
         print (option)
    tested_document =  document(path=r'Relative path',option='docx')

    try:
        #tested_document = document(path = inputfile, option = option)
        #tested_document = document(text = in_text, option = 'raw_text')
        tested_document =  document(path=r'Relative path.docx',option='docx')
    except:
        print('Failed to create document')
        return None, ' ',None


    try:
        pred, in_text,objects = tested_document.classify_text()
    except:
        print('Failed to classify document')
        return None, ' ',None
    print(pred)

    if pred == 'invoice':
        tested_document = invoice(in_text)
    elif pred == 'cv':
        tested_document = cv(in_text)
    elif pred == 'letter':
        tested_document = letter(in_text)
    else:
        print('Classifier error')

    try:
        x,text = tested_document.get_data()
    except:
        print('Failed obtain data from document')
        return None, ' ',None

    try:
        tested_document.to_json(outputfile)
    except:
        print('Json dump fail')

    return json.dumps(x)
    #return x,objects,text


if __name__ == '__main__':
     x = main(sys.argv[1:])
    #out = sort_layout_x(pages)
